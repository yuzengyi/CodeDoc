import os
import re
from docx import Document
from docx.shared import Pt
from glob import glob

class CodeDocGenerator:
    name = "软件著作权代码文档生成器"
    version = "v1.0.0"
    source_path = "E:\\project\\CodeDoc\\code"
    output_path = "E:\\project\\CodeDoc\\output"
    exclude_files = ["RuanZhuCode.iml", "README.md"]
    exclude_dirs = ["target", ".idea", "src\\test"]
    additional_files = []
    additional_dirs = []

    def __init__(self):
        self.exclude_extensions = ['mp3', 'wav', 'png', 'jpg', 'docx', 'pdf', 'exe', 'zip', 'rar']

    def generate_document(self):
        files = self.get_files_list(self.source_path)
        document_text = self.process_files(files)
        self.create_word_document(document_text)

    def get_files_list(self, root_path):
        files_list = []
        for root, dirs, files in os.walk(root_path):
            dirs[:] = [d for d in dirs if d not in self.exclude_dirs]
            for file in files:
                if not any(file.endswith(ext) for ext in self.exclude_extensions):
                    files_list.append(os.path.join(root, file))
        return files_list

    def process_files(self, files):
        text = ""
        for file in files:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
                content = re.sub("//.*", "", content)  # Remove // comments
                content = re.sub("/\*[\s\S]*?\*/", "", content)  # Remove /* */ comments
                content = re.sub("\n\s*\n", "\n", content)  # Remove empty lines
                text += content + '\n'
        return text

    def create_word_document(self, text):
        doc = Document()
        for line in text.split('\n'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)
        doc.save(os.path.join(self.output_path, f"{self.name}_{self.version}_源代码.docx"))

if __name__ == "__main__":
    generator = CodeDocGenerator()
    generator.generate_document()
